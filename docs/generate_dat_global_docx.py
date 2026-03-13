from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt


def heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def para(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style='List Bullet')


def numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style='List Number')


def codeblock(doc: Document, content: str) -> None:
    for line in content.strip('\n').split('\n'):
        p = doc.add_paragraph()
        r = p.add_run(line)
        r.font.name = 'Consolas'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
        r.font.size = Pt(9)


def cover(doc: Document) -> None:
    now = datetime.now().strftime('%Y-%m-%d %H:%M:%S')

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = p.add_run('DAT GLOBAL - Forex Multi-Agent Trading Platform')
    r.bold = True
    r.font.size = Pt(22)

    p2 = doc.add_paragraph()
    p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r2 = p2.add_run('Dossier d\'Architecture Technique Global (V1)')
    r2.font.size = Pt(14)

    p3 = doc.add_paragraph()
    p3.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r3 = p3.add_run(f'Genere automatiquement le {now}')
    r3.font.size = Pt(10)

    doc.add_paragraph('')


def build() -> Document:
    doc = Document()
    cover(doc)

    heading(doc, '0. Sommaire')
    numbered(
        doc,
        [
            'Contexte, objectifs, perimetre',
            'Architecture globale (systeme)',
            'Architecture applicative backend',
            'Architecture applicative frontend',
            'Architecture data et persistence',
            'Orchestration multi-agent et moteur de decision',
            'Execution trading, risque et garde-fous',
            'Integrations externes (Ollama, MetaApi, yfinance)',
            'Asynchrone, real-time, observabilite et audit',
            'Securite, IAM et conformite operationnelle',
            'Deploiement local, Kubernetes/Helm, CI/CD',
            'Tests, qualite, resilence et roadmap',
        ],
    )

    heading(doc, '1. Contexte, Objectifs, Perimetre')
    para(
        doc,
        'Ce document decrit l\'architecture technique globale de la plateforme IA multi-agent Forex. '
        'Il couvre l\'ensemble du produit et pas uniquement le backtest: UX, APIs, orchestration, donnees, execution, '
        'securite, observabilite, exploitation et trajectoire evolutive.',
    )
    bullets(
        doc,
        [
            'Scope fonctionnel V1: Forex uniquement (paires majeures et mineures cibles).',
            'Decisions supportees: BUY, SELL, HOLD.',
            'Modes execution: simulation, paper, live (live desactive par defaut).',
            'LLM provider principal: Ollama Cloud.',
            'Execution broker: MetaApi Cloud SDK (avec fallback REST).',
            'Contexte marche/news: Yahoo Finance via yfinance.',
            'Contraintes qualite: modularite, tracabilite, tests, observabilite, Docker Compose + preparation Kubernetes.',
        ],
    )

    heading(doc, '2. Architecture Globale (Systeme)')
    para(doc, 'Vue d\'ensemble des composants et flux entre utilisateurs, UI, API, worker et services externes.')
    heading(doc, 'Schema 1 - Vue contexte (Mermaid)', 2)
    codeblock(
        doc,
        '''
```mermaid
flowchart LR
    U1[Trader Operator] --> FE[Frontend React / Vite]
    U2[Analyst] --> FE
    U3[Admin] --> FE

    FE --> API[FastAPI Backend]
    API --> ORCH[ForexOrchestrator]
    API --> BTE[BacktestEngine]

    ORCH --> OL[Ollama Cloud]
    ORCH --> YF[yfinance]
    ORCH --> RK[RiskEngine]
    ORCH --> EX[ExecutionService]
    EX --> MA[MetaApi]

    API --> DB[(PostgreSQL)]
    API --> QD[(Qdrant)]
    API --> RB[(Redis)]
    API --> MQ[(RabbitMQ)]
    MQ --> WK[Celery Worker]

    API --> PR[/metrics Prometheus/]
    PR --> GF[Grafana]
```
''',
    )

    heading(doc, '3. Architecture Applicative Backend')
    bullets(
        doc,
        [
            'Couche API (routes): auth, runs, backtests, trading, connectors, prompts, memory, analytics, health.',
            'Couche metier (services): orchestrator, agents, risk, execution, llm, market-data, memory, prompts, trading, backtest.',
            'Couche persistence: SQLAlchemy models, migrations Alembic.',
            'Couche asynchrone: Celery + RabbitMQ + Redis.',
            'Couche observabilite: logs, metriques Prometheus, traces OpenTelemetry (optionnelle).',
        ],
    )
    heading(doc, 'Schema 2 - Modules backend (Mermaid)', 2)
    codeblock(
        doc,
        '''
```mermaid
flowchart TB
    subgraph API
      AUTH[auth.py]
      RUNS[runs.py]
      BTS[backtests.py]
      TRD[trading.py]
      CON[connectors.py]
      PRM[prompts.py]
      MEM[memory.py]
      ANA[analytics.py]
      HLT[health.py]
    end

    subgraph SVC
      ORC[services/orchestrator]
      AGT[services/orchestrator/agents]
      LLM[services/llm]
      MKT[services/market]
      EXE[services/execution]
      RSK[services/risk]
      TRA[services/trading]
      VEC[services/memory]
      BKT[services/backtest]
      PRT[services/prompts]
    end

    API --> SVC
    ORC --> AGT
    ORC --> RSK
    ORC --> EXE
    EXE --> TRA
    AGT --> LLM
    ORC --> MKT
    ORC --> VEC
    BKT --> ORC
```
''',
    )

    heading(doc, '4. Architecture Applicative Frontend')
    bullets(
        doc,
        [
            'Stack: React + TypeScript + Vite.',
            'Pages principales: login, dashboard, detail run, trading/accounts/orders/positions, connectors, prompts, backtests.',
            'Consommation API REST centralisee.',
            'Suivi de run: polling + websocket /ws/runs/{run_id}.',
            'Design dark premium, orientee lisibilite decision/risque/execution.',
        ],
    )

    heading(doc, '5. Architecture Data Et Persistence')
    para(doc, 'Persistance relationnelle + memoire vectorielle.')
    bullets(
        doc,
        [
            'PostgreSQL: runs, agent_steps, execution_orders, users, connectors, prompts, comptes MetaApi, backtests, llm logs, audit.',
            'Qdrant: index/vector store memoire long-terme.',
            'pgvector: support optionnel cote PostgreSQL (feature flag).',
            'Redis: cache/result backend et support async.',
        ],
    )

    heading(doc, 'Schema 3 - ER simplifie (Mermaid)', 2)
    codeblock(
        doc,
        '''
```mermaid
erDiagram
    USERS ||--o{ ANALYSIS_RUNS : creates
    ANALYSIS_RUNS ||--o{ AGENT_STEPS : has
    ANALYSIS_RUNS ||--o{ EXECUTION_ORDERS : has
    ANALYSIS_RUNS ||--o{ MEMORY_ENTRIES : contributes
    USERS ||--o{ BACKTEST_RUNS : creates
    BACKTEST_RUNS ||--o{ BACKTEST_TRADES : has
    USERS ||--o{ PROMPT_TEMPLATES : authors

    ANALYSIS_RUNS {
      int id PK
      string pair
      string timeframe
      string mode
      string status
      json decision
      json trace
    }
    AGENT_STEPS {
      int id PK
      int run_id FK
      string agent_name
      json input_payload
      json output_payload
    }
    EXECUTION_ORDERS {
      int id PK
      int run_id FK
      string mode
      string side
      string symbol
      float volume
      string status
    }
```
''',
    )

    heading(doc, '6. Orchestration Multi-Agent Et Moteur De Decision')
    bullets(
        doc,
        [
            'Orchestrateur central: ForexOrchestrator.',
            'Pipeline standard: technical -> news -> macro -> sentiment -> bullish -> bearish -> trader -> risk -> execution.',
            'Prompts versionnes en base pour news/bullish/bearish.',
            'Memoire vectorielle injectee dans le contexte agent.',
            'Decision finale structuree + score confiance + rationale + trace complete.',
        ],
    )

    heading(doc, 'Schema 4 - Sequence run complet (Mermaid)', 2)
    codeblock(
        doc,
        '''
```mermaid
sequenceDiagram
    participant UI
    participant API
    participant ORC as Orchestrator
    participant OL as Ollama
    participant YF as yfinance
    participant RK as Risk
    participant EX as Execution
    participant MA as MetaApi
    participant DB

    UI->>API: POST /runs
    API->>DB: create run (pending)
    API->>ORC: execute
    ORC->>YF: market + news
    ORC->>OL: news / bullish / bearish
    ORC->>RK: evaluate
    alt accepted
      ORC->>EX: execute
      EX->>MA: order
    end
    ORC->>DB: run trace + steps + order log
    API-->>UI: completed/failed
```
''',
    )

    heading(doc, '7. Execution Trading, Risque Et Garde-fous')
    bullets(
        doc,
        [
            'RiskEngine: stop loss obligatoire, controle distance SL, limite risque selon mode, sizing volume propose.',
            'ExecutionService: simulation direct, paper via MetaApi puis fallback simulation, live bloque si policy inactive.',
            'MetaApi multi-comptes: selection compte par reference + compte par defaut.',
            'Protection de production: live desactive par defaut, RBAC obligatoire sur endpoints sensibles.',
        ],
    )

    heading(doc, 'Schema 5 - Decision execution (Mermaid)', 2)
    codeblock(
        doc,
        '''
```mermaid
flowchart TD
    D[Trader decision BUY/SELL/HOLD] --> R[Risk checks]
    R -->|failed| N1[Block order]
    R -->|passed| M{Mode}
    M -->|simulation| S[Simulated execution]
    M -->|paper| P[MetaApi paper attempt]
    M -->|live| L{ALLOW_LIVE_TRADING}
    L -->|false| N2[Live blocked]
    L -->|true| V[MetaApi live order]
    P -->|provider error| F[paper_fallback simulated]
```
''',
    )

    heading(doc, '8. Integrations Externes')
    bullets(
        doc,
        [
            'Ollama Cloud: endpoint /api/chat, retries, logs de latence/tokens/cout.',
            'MetaApi: SDK principal + fallback REST, controle symbol suffix, controle tradabilite.',
            'Yahoo Finance: snapshot indicateurs + historiques + news contexte.',
        ],
    )

    heading(doc, '9. Asynchrone, Real-time, Observabilite Et Audit')
    bullets(
        doc,
        [
            'Runs: support async_execution via Celery queue analysis.',
            'Backtests: execution API (synchrone actuelle), traces metriques et logs detailles.',
            'Real-time: websocket run updates + polling front.',
            'Prometheus: analysis_runs_total, orchestrator_step_duration_seconds, llm_calls_total, llm_latency_seconds, llm_cost_usd_total, etc.',
            'Grafana: dashboard LLM/Orchestrator provisionne.',
            'Audit: journalisation des actions sensibles et des decisions.',
        ],
    )

    heading(doc, '10. Securite, IAM Et Conformite Operationnelle')
    bullets(
        doc,
        [
            'JWT + RBAC sur API.',
            'Roles: super-admin, admin, trader-operator, analyst, viewer.',
            'Validation d\'inputs via Pydantic.',
            'Gestion des secrets par variables d\'environnement.',
            'Journalisation des operations critiques (connecteurs, prompts, ordres).',
            'Separation stricte simulation/paper/live.',
        ],
    )

    heading(doc, '11. Deploiement Local, Kubernetes/Helm, CI/CD')
    bullets(
        doc,
        [
            'Local: Docker Compose (backend, worker, postgres, redis, rabbitmq, qdrant, frontend, prometheus, grafana).',
            'Kubernetes: chart Helm minimal present pour preparation deployment.',
            'CI: lint/tests backend + build frontend via GitHub Actions.',
            'Runtime default: CPU-only local.',
        ],
    )

    heading(doc, 'Schema 6 - Deploiement Compose (Mermaid)', 2)
    codeblock(
        doc,
        '''
```mermaid
flowchart LR
    FE[frontend] --> BE[backend]
    BE --> PG[(postgres)]
    BE --> R[(redis)]
    BE --> MQ[(rabbitmq)]
    BE --> Q[(qdrant)]
    BE --> O[ollama.com]
    BE --> M[metaapi]
    WK[worker] --> MQ
    WK --> PG
    WK --> R
    PROM[prometheus] --> BE
    GRAF[grafana] --> PROM
```
''',
    )

    heading(doc, '12. Strategie De Test Et Qualite')
    bullets(
        doc,
        [
            'Tests unitaires backend obligatoires (pytest).',
            'Tests integration API obligatoires.',
            'Tests e2e frontend minimaux (Playwright).',
            'Critiques de non-regression: orchestration, execution/risk, connectors, backtests.',
            'Verification de robustesse mode degrade (LLM/MetaApi/yfinance indisponibles).',
        ],
    )

    heading(doc, '13. Risques Techniques Et Recommandations')
    bullets(
        doc,
        [
            'Qdrant warning de compatibilite version client/serveur: aligner les versions.',
            'Worker Celery en root: passer en uid non-root.',
            'Backtests LLM longs sur endpoint synchrone: prevoir endpoint async + progression pour production.',
            'Runs/backtests stale: ajouter watchdog de cloture automatique.',
            'CORS OPTIONS 400 sur routes front non exposees: aligner mapping menu/endpoints.',
        ],
    )

    heading(doc, '14. Plan D\'Evolution')
    numbered(
        doc,
        [
            'Asynchroniser les backtests longs avec Celery et statut de progression.',
            'Renforcer governance prompts (approval workflow, version pinning, rollback).',
            'Ameliorer realism execution/backtest (slippage, spread, commission).',
            'Ajouter SLO/SLA observabilite et alerting.',
            'Durcir securite ops (non-root, secret manager, rotation credentials).',
            'Etendre architecture multi-environnements (dev/stage/prod) via Helm values et GitHub Actions.',
        ],
    )

    heading(doc, 'Annexe A - Inventaire Endpoints (Synthese)')
    bullets(
        doc,
        [
            'auth: /auth/login, /auth/me, /auth/bootstrap-admin',
            'runs: /runs, /runs/{id}',
            'backtests: /backtests, /backtests/{id}',
            'trading: /trading/orders, /trading/accounts, /trading/account, /trading/positions',
            'connectors: /connectors, /connectors/{name}, /connectors/{name}/test',
            'prompts: /prompts, /prompts/{id}/activate',
            'memory: /memory, /memory/search',
            'analytics: /analytics/llm-summary, /analytics/backtests-summary',
            'health: /health, /metrics, /ws/runs/{run_id}',
        ],
    )

    heading(doc, 'Annexe B - Variables Critiques')
    bullets(
        doc,
        [
            'Ollama: OLLAMA_BASE_URL, OLLAMA_API_KEY, OLLAMA_MODEL, OLLAMA_TIMEOUT_SECONDS',
            'MetaApi: METAAPI_TOKEN, METAAPI_ACCOUNT_ID, METAAPI_REGION, METAAPI_BASE_URL, METAAPI_MARKET_BASE_URL, METAAPI_SYMBOL_SUFFIX',
            'Execution controls: ALLOW_LIVE_TRADING, ENABLE_PAPER_EXECUTION',
            'Async: REDIS_URL, CELERY_BROKER_URL, CELERY_RESULT_BACKEND',
            'Memory: QDRANT_URL, QDRANT_COLLECTION, ENABLE_PGVECTOR',
            'Backtest tuning: BACKTEST_ENABLE_LLM, BACKTEST_LLM_EVERY, BACKTEST_AGENT_LOG_EVERY',
        ],
    )

    para(doc, 'Fin du DAT Global.')
    return doc


def main() -> None:
    out = Path('docs')
    out.mkdir(parents=True, exist_ok=True)
    doc = build()
    out_file = out / 'DAT_Global_Forex_MultiAgent_V1.docx'
    doc.save(out_file)
    print(out_file.resolve())


if __name__ == '__main__':
    main()
