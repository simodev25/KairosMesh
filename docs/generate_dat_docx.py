from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt


def add_title(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(22)

    p2 = doc.add_paragraph()
    p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r2 = p2.add_run(subtitle)
    r2.font.size = Pt(12)

    doc.add_paragraph('')


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style='List Bullet')


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style='List Number')


def add_code_block(doc: Document, code: str) -> None:
    for line in code.strip('\n').split('\n'):
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
        run.font.size = Pt(9)


def add_section(doc: Document, title: str, level: int = 1) -> None:
    doc.add_heading(title, level=level)


def add_paragraphs(doc: Document, lines: list[str]) -> None:
    for line in lines:
        doc.add_paragraph(line)


def build_doc() -> Document:
    doc = Document()
    generated_at = datetime.now().strftime('%Y-%m-%d %H:%M:%S')

    add_title(
        doc,
        'DAT - Dossier d\'Architecture Technique',
        f'Projet: Forex Multi-Agent Trading Platform | Version: V1 | Genere le {generated_at}',
    )

    add_section(doc, '1. Objet Du Document', 1)
    add_paragraphs(
        doc,
        [
            'Ce DAT decrit l\'architecture technique detaillee de la plateforme IA multi-agent dediee au Forex.',
            'Il couvre le backend, le frontend, l\'orchestration agents, la couche execution, la gestion du risque, la donnee, la securite, l\'observabilite, les tests et l\'exploitation.',
            'Le document est aligne avec l\'implementation actuelle du repository et les contraintes V1 (forex uniquement, paper avant live, live desactive par defaut).',
        ],
    )

    add_section(doc, '2. Perimetre Et Contexte', 1)
    add_bullets(
        doc,
        [
            'Classe d\'actifs: Forex uniquement (V1).',
            'Paires cibles V1: EURUSD, GBPUSD, USDJPY, USDCHF, AUDUSD, USDCAD, NZDUSD, EURJPY, GBPJPY, EURGBP.',
            'Timeframes V1: M5, M15, H1, H4, D1.',
            'Decisions: BUY, SELL, HOLD.',
            'Modes execution: simulation, paper, live (live protegee et desactivee par defaut).',
            'Fournisseur LLM: Ollama Cloud (obligatoire).',
            'Execution trading: MetaApi Cloud SDK + fallback REST.',
            'Donnees contexte/news: Yahoo Finance via yfinance.',
        ],
    )

    add_section(doc, '3. Vue D\'Ensemble Architecture', 1)
    add_paragraphs(
        doc,
        [
            'L\'architecture suit un decoupage modulaire: API FastAPI, worker Celery, base PostgreSQL, broker RabbitMQ, cache/result backend Redis, memoire vectorielle Qdrant/pgvector, UI React/TypeScript, stack observabilite Prometheus/Grafana.',
            'Le workflow principal est pilote par un orchestrateur central qui enchaine les agents specialises, evalue le risque puis delegue l\'execution.',
        ],
    )

    add_section(doc, 'Schema 1 - Contexte Systeme (Mermaid)', 2)
    add_code_block(
        doc,
        '''
```mermaid
flowchart LR
    U[Operateur / Analyste] --> FE[Frontend React Vite]
    FE --> API[FastAPI Backend]
    API --> ORCH[ForexOrchestrator]
    ORCH --> LLM[Ollama Cloud API]
    ORCH --> YF[yfinance / Yahoo Finance]
    ORCH --> RISK[Risk Engine]
    ORCH --> EXEC[Execution Service]
    EXEC --> META[MetaApi Cloud SDK / REST]
    API --> PG[(PostgreSQL)]
    API --> QD[(Qdrant)]
    API --> RB[(Redis)]
    API --> MQ[(RabbitMQ)]
    MQ --> WK[Celery Worker]
    API --> PROM[Prometheus /metrics]
    PROM --> GRAF[Grafana]
```
''',
    )

    add_section(doc, '4. Architecture Logique Backend', 1)
    add_bullets(
        doc,
        [
            'app/api/routes: endpoints REST (auth, runs, backtests, trading, connectors, prompts, memory, analytics, health).',
            'app/services/orchestrator: orchestration multi-agent et handoffs.',
            'app/services/llm: client Ollama Cloud avec retries, metriques et journalisation cout/latence.',
            'app/services/market: yfinance pour snapshot marche, historiques et news.',
            'app/services/trading: client MetaApi SDK + fallback REST.',
            'app/services/execution: simulation/paper/live, garde-fous, journalisation ordres.',
            'app/services/risk: validation risque, sizing, blocages.',
            'app/services/prompts: prompts versionnes en base et activation.',
            'app/services/memory: memoire vectorielle long-terme, recherche contexte.',
            'app/services/backtest: moteur backtesting (ema_rsi, agents_v1), analytics (sharpe, sortino, drawdown, profit factor).',
            'app/tasks: tache Celery pour execution asynchrone des runs.',
        ],
    )

    add_section(doc, 'Schema 2 - Composants Backend (Mermaid)', 2)
    add_code_block(
        doc,
        '''
```mermaid
flowchart TB
    subgraph API[FastAPI]
      AUTH[Auth + RBAC]
      RUNS[Runs API]
      BTS[Backtests API]
      TRD[Trading API]
      CONN[Connectors API]
      MEMAPI[Memory API]
      PRM[Prompts API]
      ANA[Analytics API]
    end

    subgraph SVC[Services]
      ORCH[Orchestrator]
      AG[Agents]
      RSK[Risk Engine]
      EXE[Execution Service]
      META[MetaApi Client]
      OLL[Ollama Client]
      YFP[YFinance Provider]
      MEM[Vector Memory]
      BTE[Backtest Engine]
      PREG[Prompt Registry]
    end

    API --> SVC
    ORCH --> AG
    ORCH --> RSK
    ORCH --> EXE
    AG --> OLL
    AG --> PREG
    ORCH --> YFP
    ORCH --> MEM
    EXE --> META
    BTE --> ORCH
```
''',
    )

    add_section(doc, '5. Workflow Multi-Agent (Run Trading)', 1)
    add_paragraphs(
        doc,
        [
            'Workflow de reference (ordre strict): technical-analyst -> news-analyst -> macro-analyst -> sentiment-agent -> bullish-researcher -> bearish-researcher -> trader-agent -> risk-manager -> execution-manager.',
            'Chaque etape est tracee en base (agent_steps) avec payload in/out.',
            'Le run global consolide decision, trace, erreurs eventuelles et statut final.',
        ],
    )

    add_section(doc, 'Schema 3 - Sequence Run (Mermaid)', 2)
    add_code_block(
        doc,
        '''
```mermaid
sequenceDiagram
    participant UI as Frontend
    participant API as FastAPI /runs
    participant OR as ForexOrchestrator
    participant YF as yfinance
    participant OL as Ollama
    participant RK as RiskEngine
    participant EX as ExecutionService
    participant MT as MetaApi
    participant DB as PostgreSQL

    UI->>API: POST /runs
    API->>DB: create analysis_run(pending)
    API->>OR: execute(context)
    OR->>YF: market snapshot + news
    OR->>OL: news analyst
    OR->>OL: bullish researcher
    OR->>OL: bearish researcher
    OR->>RK: evaluate risk
    alt accepted and decision in BUY/SELL
      OR->>EX: execute order
      EX->>MT: place order (paper/live)
      EX->>DB: execution_order
    else HOLD or rejected
      OR->>DB: skipped execution
    end
    OR->>DB: agent_steps + run trace + decision
    API-->>UI: run completed/failed
```
''',
    )

    add_section(doc, '6. Workflow Backtest', 1)
    add_bullets(
        doc,
        [
            'Strategie legacy: ema_rsi (deterministe).',
            'Strategie cible: agents_v1 (pipeline partage avec orchestrateur).',
            'Source de workflow exposee en metriques: workflow_source=ForexOrchestrator.analyze_context.',
            'LLM en backtest: configurable via BACKTEST_ENABLE_LLM et BACKTEST_LLM_EVERY.',
            'Execution reelle desactivee en backtest (execution_mode=disabled-in-backtest).',
        ],
    )

    add_section(doc, 'Schema 4 - Sequence Backtest (Mermaid)', 2)
    add_code_block(
        doc,
        '''
```mermaid
sequenceDiagram
    participant UI as Frontend Backtests
    participant API as FastAPI /backtests
    participant BTE as BacktestEngine
    participant OR as Orchestrator.analyze_context
    participant OL as Ollama (optionnel)
    participant DB as PostgreSQL

    UI->>API: POST /backtests(strategy=agents_v1)
    API->>DB: create backtest_run(running)
    API->>BTE: run(pair, timeframe, dates, db)
    loop Candles preparees
      BTE->>OR: analyze_context(mode=backtest)
      alt BACKTEST_ENABLE_LLM=true and cadence respectee
        OR->>OL: appels LLM (news/debate)
      end
      BTE->>BTE: signal BUY/SELL/HOLD
    end
    BTE->>DB: backtest_trades + metrics + equity_curve
    API-->>UI: status completed/failed
```
''',
    )

    add_section(doc, '7. Donnees Et Modele Persistant', 1)
    add_paragraphs(
        doc,
        [
            'Base principale: PostgreSQL (SQLAlchemy + Alembic).',
            'Tables coeur execution: analysis_runs, agent_steps, execution_orders.',
            'Tables administration: users, connector_configs, metaapi_accounts, prompt_templates, audit_logs.',
            'Tables data science: memory_entries (embedding), llm_call_logs, backtest_runs, backtest_trades.',
            'Memoire vectorielle: Qdrant principal, pgvector en option (ENABLE_PGVECTOR).',
        ],
    )

    add_section(doc, 'Schema 5 - Modele Relationnel Simplifie (Mermaid)', 2)
    add_code_block(
        doc,
        '''
```mermaid
erDiagram
    USERS ||--o{ ANALYSIS_RUNS : creates
    ANALYSIS_RUNS ||--o{ AGENT_STEPS : contains
    ANALYSIS_RUNS ||--o{ EXECUTION_ORDERS : produces
    ANALYSIS_RUNS ||--o{ MEMORY_ENTRIES : enriches
    USERS ||--o{ BACKTEST_RUNS : creates
    BACKTEST_RUNS ||--o{ BACKTEST_TRADES : contains
    USERS ||--o{ PROMPT_TEMPLATES : authors

    USERS {
      int id PK
      string email
      string role
      bool is_active
    }
    ANALYSIS_RUNS {
      int id PK
      string pair
      string timeframe
      string mode
      string status
      json decision
      json trace
      text error
      int created_by_id FK
    }
    AGENT_STEPS {
      int id PK
      int run_id FK
      string agent_name
      string status
      json input_payload
      json output_payload
    }
    EXECUTION_ORDERS {
      int id PK
      int run_id FK
      string mode
      string side
      string symbol
      float volume
      string status
      json request_payload
      json response_payload
    }
    MEMORY_ENTRIES {
      int id PK
      string pair
      string timeframe
      string source_type
      text summary
      vector embedding
      json payload
      int run_id FK
    }
```
''',
    )

    add_section(doc, '8. API Et Contrats', 1)
    add_paragraphs(doc, ['Principaux endpoints observes dans le code:'])
    add_bullets(
        doc,
        [
            'Auth: POST /api/v1/auth/login, GET /api/v1/auth/me, POST /api/v1/auth/bootstrap-admin',
            'Runs: GET /api/v1/runs, POST /api/v1/runs, GET /api/v1/runs/{id}',
            'Backtests: GET /api/v1/backtests, POST /api/v1/backtests, GET /api/v1/backtests/{id}',
            'Trading: GET /api/v1/trading/orders, GET/POST/PATCH /api/v1/trading/accounts, GET /api/v1/trading/account, GET /api/v1/trading/positions',
            'Connectors: GET /api/v1/connectors, PUT /api/v1/connectors/{name}, POST /api/v1/connectors/{name}/test',
            'Prompts: GET/POST /api/v1/prompts, POST /api/v1/prompts/{id}/activate',
            'Memory: GET /api/v1/memory, POST /api/v1/memory/search',
            'Analytics: GET /api/v1/analytics/llm-summary, GET /api/v1/analytics/backtests-summary',
            'Health: GET /api/v1/health, GET /metrics, GET / (root), WS /ws/runs/{run_id}',
        ],
    )

    add_section(doc, '9. Securite', 1)
    add_bullets(
        doc,
        [
            'Authentification JWT avec roles: super-admin, admin, trader-operator, analyst, viewer.',
            'RBAC applique aux endpoints sensibles (trading/accounts/connectors/prompts/backtests/runs).',
            'Live trading bloque par defaut (ALLOW_LIVE_TRADING=false).',
            'Paper trading controle (ENABLE_PAPER_EXECUTION=true/false).',
            'Validation d\'inputs (Pydantic) et garde-fous risque obligatoires avant execution.',
            'Masquage/separation des secrets via variables d\'environnement.',
        ],
    )

    add_section(doc, 'Schema 6 - Gouvernance Des Modes D\'Execution', 2)
    add_code_block(
        doc,
        '''
```mermaid
flowchart TD
    D[Decision trader BUY/SELL/HOLD] --> R[Risk Engine]
    R -->|Rejected| X[No Trade]
    R -->|Accepted| M{Mode}
    M -->|simulation| S[Simulated Fill]
    M -->|paper| P[MetaApi Paper Attempt]
    M -->|live| L{ALLOW_LIVE_TRADING}
    L -->|false| LB[Blocked]
    L -->|true| LV[MetaApi Live Attempt]
    P -->|MetaApi unavailable| PF[Paper Fallback Simulated]
```
''',
    )

    add_section(doc, '10. Observabilite Et Audit', 1)
    add_bullets(
        doc,
        [
            'Metriques Prometheus exposees via /metrics.',
            'Dashboard Grafana provisionne (latence/couts LLM, orchestrator).',
            'Logs applicatifs structures (runs, backtests, agents, execution, llm).',
            'Persistance des appels LLM (llm_call_logs) avec tokens, cout USD, latence, statut.',
            'Traçabilite des decisions agents en base (agent_steps + analysis_runs.trace).',
        ],
    )

    add_section(doc, '11. Deploiement Et Exploitation', 1)
    add_paragraphs(doc, ['Execution locale standard via Docker Compose.'])
    add_section(doc, 'Schema 7 - Deploiement Docker Compose (Mermaid)', 2)
    add_code_block(
        doc,
        '''
```mermaid
flowchart LR
    FE[frontend:5173] --> BE[backend:8000]
    BE --> PG[(postgres:5432)]
    BE --> RD[(redis:6379)]
    BE --> MQ[(rabbitmq:5672)]
    BE --> QD[(qdrant:6333)]
    BE --> OC[ollama.com]
    BE --> MA[MetaApi]
    WK[worker] --> MQ
    WK --> PG
    WK --> RD
    PR[prometheus:9090] --> BE
    GF[grafana:3000] --> PR
```
''',
    )

    add_section(doc, '12. Configuration Critique (.env)', 1)
    add_bullets(
        doc,
        [
            'LLM: OLLAMA_BASE_URL, OLLAMA_API_KEY, OLLAMA_MODEL, OLLAMA_TIMEOUT_SECONDS.',
            'MetaApi: METAAPI_TOKEN, METAAPI_ACCOUNT_ID, METAAPI_REGION, METAAPI_BASE_URL, METAAPI_MARKET_BASE_URL, METAAPI_SYMBOL_SUFFIX.',
            'Execution controls: ALLOW_LIVE_TRADING, ENABLE_PAPER_EXECUTION.',
            'Queues: REDIS_URL, CELERY_BROKER_URL, CELERY_RESULT_BACKEND, CELERY_IGNORE_RESULT.',
            'Memory: QDRANT_URL, QDRANT_COLLECTION, MEMORY_VECTOR_SIZE, ENABLE_PGVECTOR.',
            'Backtest logging/LLM: LOG_AGENT_STEPS, BACKTEST_AGENT_LOG_EVERY, BACKTEST_ENABLE_LLM, BACKTEST_LLM_EVERY.',
        ],
    )

    add_section(doc, '13. Tests Et Qualite', 1)
    add_bullets(
        doc,
        [
            'Unit tests backend: pytest (moteur backtest, agents, trader, API runs/backtests).',
            'Integration API: endpoints runs/backtests/trading/connectors.',
            'Frontend: build + e2e minimal Playwright.',
            'Objectif qualite: gestion explicite des erreurs, entree validee, resilience en mode degrade.',
        ],
    )

    add_section(doc, '14. Modes Degrades Et Resilience', 1)
    add_bullets(
        doc,
        [
            'Ollama indisponible/auth KO: fallback deterministe + erreur explicite.',
            'MetaApi indisponible: fallback paper-simulated en mode paper; blocage en live si policy non satisfaite.',
            'yfinance indisponible: snapshots/news degrades, run possible avec contexte partiel.',
            'Qdrant indisponible: recherche memoire degradee, run non bloque.',
        ],
    )

    add_section(doc, '15. Ecarts Techniques Observes Et Recommandations', 1)
    add_bullets(
        doc,
        [
            'Compatibilite Qdrant: client 1.15.1 vs serveur 1.13.2 (warning). Recommandation: aligner versions.',
            'Worker Celery en root (warning securite). Recommandation: executer avec uid non-root.',
            'Backtests LLM longs peuvent saturer une requete HTTP synchrone. Recommandation: endpoint async de backtest avec polling/ws.',
            'CORS preflight 400 sur certains endpoints front non exposes. Recommandation: harmoniser navigation/menu avec routes disponibles.',
        ],
    )

    add_section(doc, '16. Roadmap Technique Proposee (V1.1 -> V2)', 1)
    add_numbered(
        doc,
        [
            'Asynchroniser les backtests longs via Celery + statut progressif.',
            'Ajouter un watchdog qui cloture automatiquement les runs/backtests stale.',
            'Renforcer la gouvernance prompts (A/B testing, rollback, approvals).',
            'Ajouter simulation de slippage/spread/commission plus realiste.',
            'Mettre en place CI quality gates plus stricts (coverage, lint, type checks).',
            'Durcir la securite ops (non-root, rotation secrets, audit events complets).',
        ],
    )

    add_section(doc, '17. Annexes - Schemas Complementaires', 1)
    add_section(doc, 'Schema 8 - Machine A Etats Analysis Run (Mermaid)', 2)
    add_code_block(
        doc,
        '''
```mermaid
stateDiagram-v2
    [*] --> pending
    pending --> queued
    queued --> running
    pending --> running
    running --> completed
    running --> failed
```
''',
    )

    add_section(doc, 'Schema 9 - Machine A Etats Backtest Run (Mermaid)', 2)
    add_code_block(
        doc,
        '''
```mermaid
stateDiagram-v2
    [*] --> running
    running --> completed
    running --> failed
```
''',
    )

    add_paragraphs(
        doc,
        [
            'Fin du document DAT.',
            'Note: les schemas Mermaid sont fournis au format texte pour versionning et revue technique. Ils peuvent etre rendus visuellement dans un outil compatible Mermaid.',
        ],
    )

    return doc


def main() -> None:
    out_dir = Path('docs')
    out_dir.mkdir(parents=True, exist_ok=True)
    doc = build_doc()
    out_path = out_dir / 'DAT_Forex_MultiAgent_V1.docx'
    doc.save(out_path)
    print(str(out_path.resolve()))


if __name__ == '__main__':
    main()
